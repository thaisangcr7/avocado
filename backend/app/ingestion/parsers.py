"""Type-specific parsers.

Each returns a `ParsedDocument`: text blocks for retrieval, plus — for
spreadsheets — structured tables for the analysis engine. Spreadsheets produce
*both*, which is the point of §7: the embedded text answers "which file
mentions Q3 revenue", the table is what pandas computes against.

Parsing is CPU-bound and synchronous underneath (pypdf, openpyxl, pandas), so
callers run these in a worker thread.
"""

from __future__ import annotations

import csv
import io
from typing import Any

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.errors import ValidationError
from app.ingestion.types import ParsedColumn, ParsedDocument, ParsedTable, TextBlock
from app.models.enums import DocumentType

# Guards against a pathological upload turning into an unbounded parse.
MAX_TABLE_ROWS = 200_000
MAX_SAMPLE_VALUES = 5


def _describe_frame(frame: pd.DataFrame) -> list[ParsedColumn]:
    columns: list[ParsedColumn] = []
    for name in frame.columns:
        series = frame[name]
        samples = [
            None if pd.isna(v) else (v.item() if hasattr(v, "item") else v)
            for v in series.dropna().head(MAX_SAMPLE_VALUES).tolist()
        ]
        columns.append(
            ParsedColumn(
                name=str(name),
                dtype=str(series.dtype),
                null_count=int(series.isna().sum()),
                sample_values=[str(s) if s is not None else None for s in samples],
            )
        )
    return columns


def _frame_to_text(frame: pd.DataFrame, table_name: str) -> str:
    """A compact textual view of a table, for retrieval.

    The head of the data plus the schema is enough for a semantic match; the
    full table lives in object storage for the analysis engine.
    """
    schema = ", ".join(f"{c} ({frame[c].dtype})" for c in frame.columns)
    preview = frame.head(20).to_csv(index=False)
    return (
        f"Table: {table_name}\n"
        f"Rows: {len(frame)}, Columns: {len(frame.columns)}\n"
        f"Schema: {schema}\n\n"
        f"First rows:\n{preview}"
    )


def _normalise_frame(frame: pd.DataFrame) -> pd.DataFrame:
    if len(frame) > MAX_TABLE_ROWS:
        frame = frame.head(MAX_TABLE_ROWS)
    # Unnamed columns come from spreadsheets with blank header cells; giving
    # them stable names keeps generated code referable.
    frame.columns = [
        str(c) if not str(c).startswith("Unnamed:") else f"column_{i}"
        for i, c in enumerate(frame.columns)
    ]
    return frame


def _table_from_frame(frame: pd.DataFrame, name: str, index: int) -> ParsedTable:
    frame = _normalise_frame(frame)
    buffer = io.StringIO()
    frame.to_csv(buffer, index=False)
    return ParsedTable(
        name=name,
        sheet_index=index,
        columns=_describe_frame(frame),
        row_count=len(frame),
        csv_bytes=buffer.getvalue().encode("utf-8"),
    )


def parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"Could not read PDF '{filename}'.") from exc

    blocks: list[TextBlock] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            # One unreadable page shouldn't discard the rest of the document.
            text = ""
        if text.strip():
            blocks.append(TextBlock(content=text, metadata={"page": page_number}))

    return ParsedDocument(
        text_blocks=blocks,
        page_count=len(reader.pages),
        metadata={
            "parser": "pypdf",
            "pages_with_text": len(blocks),
            # A PDF with pages but no extractable text is almost always a scan;
            # flagging it is what lets the OCR/vision path pick it up.
            "likely_scanned": len(reader.pages) > 0 and len(blocks) == 0,
        },
    )


def parse_docx(data: bytes, filename: str) -> ParsedDocument:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"Could not read Word document '{filename}'.") from exc

    blocks: list[TextBlock] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            blocks.append(
                TextBlock(
                    content="\n".join(buffer),
                    metadata={"section": current_heading} if current_heading else {},
                )
            )
            buffer.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # Headings become section boundaries so citations can name a section.
        if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
            flush()
            current_heading = text
            buffer.append(text)
        else:
            buffer.append(text)
    flush()

    tables: list[ParsedTable] = []
    for index, table in enumerate(document.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        frame = pd.DataFrame(rows[1:], columns=rows[0])
        tables.append(_table_from_frame(frame, f"table_{index + 1}", index))
        blocks.append(
            TextBlock(
                content=_frame_to_text(frame, f"table_{index + 1}"),
                metadata={"table": f"table_{index + 1}"},
            )
        )

    return ParsedDocument(
        text_blocks=blocks,
        tables=tables,
        metadata={"parser": "python-docx", "table_count": len(tables)},
    )


def parse_xlsx(data: bytes, filename: str) -> ParsedDocument:
    try:
        sheets: dict[str, pd.DataFrame] = pd.read_excel(
            io.BytesIO(data), sheet_name=None, engine="openpyxl"
        )
    except Exception as exc:
        raise ValidationError(f"Could not read spreadsheet '{filename}'.") from exc

    blocks: list[TextBlock] = []
    tables: list[ParsedTable] = []
    for index, (sheet_name, frame) in enumerate(sheets.items()):
        if frame.empty:
            continue
        table = _table_from_frame(frame, sheet_name, index)
        tables.append(table)
        blocks.append(
            TextBlock(
                content=_frame_to_text(_normalise_frame(frame), sheet_name),
                metadata={"sheet": sheet_name, "sheet_index": index},
            )
        )

    return ParsedDocument(
        text_blocks=blocks,
        tables=tables,
        metadata={
            "parser": "pandas/openpyxl",
            "sheet_names": list(sheets.keys()),
            "sheet_count": len(tables),
        },
    )


def _sniff_delimiter(sample: str) -> str:
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except csv.Error:
        return ","


def parse_csv(data: bytes, filename: str) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    delimiter = _sniff_delimiter(text[:8192])
    try:
        frame = pd.read_csv(io.StringIO(text), sep=delimiter)
    except Exception as exc:
        raise ValidationError(f"Could not read CSV '{filename}'.") from exc

    if frame.empty:
        raise ValidationError(f"CSV '{filename}' contains no rows.")

    name = filename.rsplit(".", 1)[0]
    table = _table_from_frame(frame, name, 0)
    return ParsedDocument(
        text_blocks=[
            TextBlock(
                content=_frame_to_text(_normalise_frame(frame), name),
                metadata={"table": name},
            )
        ],
        tables=[table],
        metadata={"parser": "pandas", "delimiter": delimiter},
    )


def parse_text(data: bytes, filename: str) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    if not text.strip():
        raise ValidationError(f"'{filename}' contains no text.")
    return ParsedDocument(
        text_blocks=[TextBlock(content=text)],
        metadata={"parser": "plain-text"},
    )


def parse_markdown(data: bytes, filename: str) -> ParsedDocument:
    """Markdown, split on headings so sections stay citable."""
    text = data.decode("utf-8", errors="replace")
    if not text.strip():
        raise ValidationError(f"'{filename}' contains no text.")

    blocks: list[TextBlock] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            blocks.append(
                TextBlock(
                    content=body,
                    metadata={"section": current_heading} if current_heading else {},
                )
            )
        buffer.clear()

    for line in text.splitlines():
        if line.startswith("#"):
            flush()
            current_heading = line.lstrip("#").strip()
        buffer.append(line)
    flush()

    return ParsedDocument(
        text_blocks=blocks or [TextBlock(content=text)],
        metadata={"parser": "markdown"},
    )


SYNC_PARSERS: dict[DocumentType, Any] = {
    DocumentType.PDF: parse_pdf,
    DocumentType.DOCX: parse_docx,
    DocumentType.XLSX: parse_xlsx,
    DocumentType.CSV: parse_csv,
    DocumentType.TEXT: parse_text,
    DocumentType.MARKDOWN: parse_markdown,
    # A voice document's stored object is its transcript, so it parses as text.
    # The audio stays referenced in metadata rather than being what is parsed.
    DocumentType.AUDIO: parse_text,
}
