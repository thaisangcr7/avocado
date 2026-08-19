"""Type-specific parsers, exercised on real files built in-memory."""

from __future__ import annotations

import io

import openpyxl
import pytest
from docx import Document as DocxDocument

from app.core.errors import ValidationError
from app.ingestion.parsers import (
    parse_csv,
    parse_docx,
    parse_markdown,
    parse_text,
    parse_xlsx,
)


def build_xlsx(sheets: dict[str, list[list]]) -> bytes:
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    for name, rows in sheets.items():
        sheet = workbook.create_sheet(name)
        for row in rows:
            sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def build_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    document = DocxDocument()
    for style, text in paragraphs:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- CSV -------------------------------------------------------------------


def test_csv_yields_both_text_and_a_structured_table():
    data = b"region,revenue\nNorth,100\nSouth,200\n"
    parsed = parse_csv(data, "sales.csv")

    assert parsed.text_blocks  # retrievable
    assert len(parsed.tables) == 1  # analysable

    table = parsed.tables[0]
    assert table.row_count == 2
    assert [c.name for c in table.columns] == ["region", "revenue"]
    assert b"North" in table.csv_bytes


def test_csv_column_types_and_samples_are_described():
    parsed = parse_csv(b"name,amount\nA,10\nB,20\n", "x.csv")
    columns = {c.name: c for c in parsed.tables[0].columns}
    assert "int" in columns["amount"].dtype
    assert columns["name"].sample_values[:2] == ["A", "B"]


def test_csv_null_counts_are_recorded():
    parsed = parse_csv(b"a,b\n1,\n2,x\n", "x.csv")
    columns = {c.name: c for c in parsed.tables[0].columns}
    assert columns["b"].null_count == 1


def test_semicolon_delimited_csv_is_detected():
    parsed = parse_csv(b"a;b;c\n1;2;3\n4;5;6\n", "euro.csv")
    assert parsed.metadata["delimiter"] == ";"
    assert [c.name for c in parsed.tables[0].columns] == ["a", "b", "c"]


def test_an_empty_csv_is_rejected():
    with pytest.raises(ValidationError):
        parse_csv(b"a,b\n", "empty.csv")


# --- XLSX ------------------------------------------------------------------


def test_every_sheet_becomes_its_own_table():
    data = build_xlsx(
        {
            "Revenue": [["region", "amount"], ["North", 100], ["South", 200]],
            "Costs": [["team", "spend"], ["Eng", 50]],
        }
    )
    parsed = parse_xlsx(data, "book.xlsx")

    assert len(parsed.tables) == 2
    assert {t.name for t in parsed.tables} == {"Revenue", "Costs"}
    assert parsed.metadata["sheet_count"] == 2
    # Each sheet is separately citable.
    assert {b.metadata["sheet"] for b in parsed.text_blocks} == {"Revenue", "Costs"}


def test_empty_sheets_are_skipped():
    data = build_xlsx({"Data": [["a"], [1]], "Blank": []})
    parsed = parse_xlsx(data, "book.xlsx")
    assert [t.name for t in parsed.tables] == ["Data"]


def test_blank_header_cells_get_stable_column_names():
    """Generated code needs a name it can actually reference."""
    data = build_xlsx({"S": [["region", None], ["North", 5]]})
    parsed = parse_xlsx(data, "book.xlsx")
    names = [c.name for c in parsed.tables[0].columns]
    assert all(not n.startswith("Unnamed") for n in names)


def test_a_corrupt_spreadsheet_is_rejected_clearly():
    with pytest.raises(ValidationError, match="Could not read spreadsheet"):
        parse_xlsx(b"PK\x03\x04 not really a workbook", "broken.xlsx")


# --- DOCX ------------------------------------------------------------------


def test_docx_headings_become_section_metadata():
    data = build_docx(
        [
            ("Heading 1", "Expenses Policy"),
            ("", "Submit receipts within thirty days."),
            ("Heading 1", "Travel Policy"),
            ("", "Book flights two weeks ahead."),
        ]
    )
    parsed = parse_docx(data, "handbook.docx")

    sections = {b.metadata.get("section") for b in parsed.text_blocks}
    assert "Expenses Policy" in sections
    assert "Travel Policy" in sections


def test_docx_body_text_is_extracted():
    data = build_docx([("", "The quick brown fox jumps over the lazy dog.")])
    parsed = parse_docx(data, "doc.docx")
    assert "quick brown fox" in parsed.text_blocks[0].content


def test_a_corrupt_word_file_is_rejected_clearly():
    with pytest.raises(ValidationError, match="Could not read Word document"):
        parse_docx(b"not a docx at all", "broken.docx")


# --- Text and Markdown -----------------------------------------------------


def test_plain_text_is_parsed():
    parsed = parse_text(b"Hello there, this is a note.", "note.txt")
    assert parsed.text_blocks[0].content == "Hello there, this is a note."


def test_empty_text_is_rejected():
    with pytest.raises(ValidationError):
        parse_text(b"   \n  ", "blank.txt")


def test_markdown_splits_on_headings():
    content = b"# Intro\n\nWelcome.\n\n## Details\n\nMore here.\n"
    parsed = parse_markdown(content, "doc.md")
    sections = [b.metadata.get("section") for b in parsed.text_blocks]
    assert "Intro" in sections
    assert "Details" in sections


def test_markdown_without_headings_still_parses():
    parsed = parse_markdown(b"Just a paragraph with no heading.", "flat.md")
    assert parsed.text_blocks
    assert "Just a paragraph" in parsed.text_blocks[0].content


def test_invalid_utf8_is_replaced_rather_than_crashing():
    parsed = parse_text(b"valid text \xff\xfe more text", "weird.txt")
    assert "valid text" in parsed.text_blocks[0].content
